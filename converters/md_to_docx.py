#!/usr/bin/env python3

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_md():
    # Create a new Document
    doc = Document()
    
    # Set the margins (in inches)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read the Markdown file
    with open('business/mass_ai_grant.md', 'r') as file:
        md_lines = file.readlines()
    
    # Process the Markdown content
    i = 0
    while i < len(md_lines):
        line = md_lines[i].strip()
        
        # Title (# heading)
        if line.startswith('# '):
            p = doc.add_heading(line[2:], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Section heading (## heading)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        
        # Subsection heading (### heading)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        
        # Bold text (**text**)
        elif line.startswith('**') and line.endswith('**') and ':' in line:
            # This is for the organization names in the Consortium Qualifications section
            parts = line[2:-2].split(':')
            p = doc.add_paragraph()
            p.add_run(parts[0] + ':').bold = True
            if len(parts) > 1:
                p.add_run(parts[1])
        
        # Bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Numbered list
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
            # Extract the number
            num_end = line.find('.')
            num = int(line[:num_end])
            
            # Extract the content after the number and period
            content = line[num_end+2:]
            
            # Check if the content has bold parts
            if '**' in content:
                p = doc.add_paragraph(style='List Number')
                parts = content.split('**')
                
                for j, part in enumerate(parts):
                    if j % 2 == 1:  # Odd indices are the bold parts
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
            else:
                doc.add_paragraph(content, style='List Number')
        
        # Regular paragraph
        elif line and not line.startswith('[') and not line.startswith('('):
            doc.add_paragraph(line)
        
        # Handle empty lines
        elif not line:
            doc.add_paragraph()
        
        i += 1
    
    # Save the document
    output_path = 'business/output/Massachusetts_AI_Models_Challenge_Application.docx'
    doc.save(output_path)
    print(f"Converted Markdown to DOCX: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    try:
        create_docx_from_md()
    except Exception as e:
        print(f"Error: {e}")
        print("You may need to install python-docx. Run: pip3 install python-docx") 